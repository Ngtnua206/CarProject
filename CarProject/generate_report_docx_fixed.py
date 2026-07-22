from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

output_path = r"D:\Code\Code\WebMVC\CarProject\BaoCao_HeThong_MyLxCar.docx"

def set_style(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    for row_values in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_values):
            row_cells[i].text = value
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
set_style(doc)

# Cover page
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.add_run('BÁO CÁO HỆ THỐNG MYLXCAR')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Arial'

p2 = doc.add_paragraph()
p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run2 = p2.add_run('Website bán xe hạng sang, quản lý showroom, đơn cọc, lịch hẹn, thống kê doanh thu và tích hợp thanh toán')
run2.italic = True
run2.font.size = Pt(12)

doc.add_paragraph('Sinh viên: Nhóm phát triển MyLxCar')
doc.add_paragraph('Môn học: Hệ thống thông tin quản lý')
doc.add_paragraph('Ngày tạo: 22/07/2026')
doc.add_paragraph('Nền tảng: ASP.NET Core Razor Pages + SQL Server + Docker + Cloudflare + Tailscale')
add_page_break(doc)

# Content
add_paragraph(doc, '1. Tổng quan hệ thống', bold=True)
doc.add_heading('1. Tổng quan hệ thống', level=1)
add_paragraph(doc, 'Hệ thống MyLxCar là nền tảng web bán hàng và quản trị showroom xe hạng sang. Từ việc đọc toàn bộ code, hệ thống gồm các thành phần chính: giao diện Razor Pages, service layer, Entity Framework Core, SQL Server chạy trong Docker, và middleware xác thực/ phân quyền.')
add_paragraph(doc, 'Hệ thống hỗ trợ người dùng xem danh sách xe, xem chi tiết, so sánh xe, đặt lịch lái thử, đặt cọc, quản lý giỏ hàng, đăng ký tài khoản, xác nhận email, thanh toán qua Sepay, và quản trị showroom cho Admin/Quản lý.')
add_table(doc, ['Thành phần', 'Vai trò'], [
    ['Frontend', 'Razor Pages + Bootstrap + Chart.js'],
    ['Backend', 'ASP.NET Core 10, middleware, service classes'],
    ['Database', 'SQL Server 2022 trên Docker'],
    ['Bảo mật', 'Session, JWT, Google OAuth, role-based authorization'],
    ['Thanh toán', 'Sepay QR + webhook'],
    ['Giám sát', 'Serilog, activity log, audit logs']
])
add_page_break(doc)

add_paragraph(doc, '2. Kiến trúc tổng thể và luồng xử lý', bold=True)
doc.add_heading('2. Kiến trúc tổng thể và luồng xử lý', level=1)
add_paragraph(doc, 'Ứng dụng được xây dựng theo kiến trúc 3 lớp: Presentation, Application/Service, Data Access. Presentation là Razor Pages và Controllers. Service layer chứa logic về giỏ hàng, email, QR thanh toán, JWT và log. Data layer là AppDbContext và các entity models.')
add_bullets(doc, [
    'Luồng đăng nhập: request -> middleware -> authentication -> page handler -> service -> database',
    'Luồng đặt cọc: giỏ hàng -> tạo đơn cọc -> gọi Sepay -> lưu giao dịch -> cập nhật trạng thái',
    'Luồng quản trị: Admin/Quản lý -> dashboard -> CRUD -> ghi log hệ thống',
    'Luồng lịch hẹn lái thử: user -> form -> lưu vào LichHenLaiThu -> admin xem và xử lý'
])
add_paragraph(doc, '[CHÈN ẢNH: Sơ đồ kiến trúc tổng thể (frontend, backend, database, Docker, Cloudflare, Tailscale)]')
add_page_break(doc)

add_paragraph(doc, '3. Cấu trúc code và logic nghiệp vụ của hệ thống', bold=True)
doc.add_heading('3. Cấu trúc code và logic nghiệp vụ của hệ thống', level=1)
add_paragraph(doc, 'Program.cs là điểm khởi đầu của ứng dụng. Ở đây, hệ thống đăng ký authentication, session, JWT, Google OAuth, DbContext, service scoped và middleware phân quyền theo đường dẫn. Đây là nơi quy định mọi route /Admin/* chỉ dành cho Admin, còn /Orders/* và /QuanLy/* yêu cầu đăng nhập.')
add_paragraph(doc, 'AppDbContext định nghĩa các DbSet và các mối quan hệ giữa bảng. Trong code, các bảng như HangXe, DongXe, PhienBanXe, TaiKhoan, DonDatCoc, HoaDonMuaXe, LichHenLaiThu, ChiNhanhShowroom, ChuongTrinhKhuyenMai, QuangCaoBanner, KenhTuVan, NhatKyHeThong, ThongKeTongHop_Boss và GioHang đều được ánh xạ với các khóa ngoại.')
add_table(doc, ['Module', 'Logic chính'], [
    ['CartService', 'Quản lý giỏ hàng, số lượng, tổng tiền cọc, kiểm tra tối thiểu 3 xe'],
    ['EmailService', 'Gửi email xác nhận đăng ký sử dụng SMTP'],
    ['SepayService', 'Tạo QR và verify webhook thanh toán'],
    ['DbInitializer', 'Seed dữ liệu ban đầu và tạo user AppReader'],
    ['ActivityLogService', 'Ghi log hành động người dùng, trình duyệt, IP, route']
])
add_page_break(doc)

add_paragraph(doc, '4. Cơ sở dữ liệu và mô hình dữ liệu', bold=True)
doc.add_heading('4. Cơ sở dữ liệu và mô hình dữ liệu', level=1)
add_paragraph(doc, 'Cơ sở dữ liệu được thiết kế trên SQL Server 2022 chạy trong container Docker. Connection string hiện tại dùng Server=100.108.48.1,1433;Database=CarShopDb;User Id=sa;Password=Iumaioanhh@2024;TrustServerCertificate=True;Encrypt=False;.')
add_bullets(doc, [
    'Database Name: CarShopDb',
    'Container SQL: sqlserver-carproject',
    'Port nội bộ: 1433',
    'Phương thức kết nối: SQL Server Authentication',
    'Mục tiêu: nhóm có thể truy cập DB từ xa mà không cần public IP phơi lộ'
])
add_paragraph(doc, '[CHÈN ẢNH: Màn hình kết nối SQL Server Management Studio / Azure Data Studio vào database]')
add_page_break(doc)

add_paragraph(doc, '5. Phân tích chuẩn hóa dữ liệu 3NF', bold=True)
doc.add_heading('5. Phân tích chuẩn hóa dữ liệu 3NF', level=1)
add_paragraph(doc, 'Mô hình dữ liệu hiện tại đã có hướng đi đúng đắn về phân tách thực thể và khóa ngoại, nhưng để đạt chuẩn 3NF một cách thỏa đáng hơn, cần lưu ý một số điểm.')
add_table(doc, ['Bảng', 'Đánh giá 3NF', 'Gợi ý cải tiến'], [
    ['HangXe', 'Đạt', 'Giữ nguyên'],
    ['DongXe', 'Đạt', 'Giữ nguyên và nối với HangXe'],
    ['PhienBanXe', 'Đạt', 'Giữ nguyên, phân loại theo dong xe'],
    ['TaiKhoan', 'Đạt', 'Giữ thông tin đăng nhập và vai trò'],
    ['DonDatCoc', 'Gần 3NF', 'Tách thông tin khách hàng sang bảng khách hàng hoặc bảng chi tiết'],
    ['HoaDonMuaXe', 'Đạt', 'Không nên lưu trùng các trường khác ngoài thông tin hóa đơn'],
    ['LichHenLaiThu', 'Đạt', 'Có thể tách trạng thái sang bảng tham chiếu trạng thái'],
    ['GioHang', 'Đạt', 'Phụ thuộc vào tài khoản và phiên bản xe']
])
add_paragraph(doc, 'Kết luận: mô hình hiện tại có thể nhìn là gần 3NF, phù hợp cho học tập và triển khai ban đầu.')
add_page_break(doc)

add_paragraph(doc, '6. Mối quan hệ giữa các bảng và ERD', bold=True)
doc.add_heading('6. Mối quan hệ giữa các bảng và ERD', level=1)
add_paragraph(doc, 'Hệ thống dùng nhiều bảng liên kết với nhau bằng khóa ngoại. Quan hệ chính: HangXe 1-n DongXe; DongXe 1-n PhienBanXe; TaiKhoan 1-n DonDatCoc; PhienBanXe 1-n DonDatCoc; DonDatCoc 1-1 HoaDonMuaXe; TaiKhoan 1-n LichHenLaiThu; DongXe 1-n LichHenLaiThu; ChiNhanhShowroom 1-n LichHenLaiThu.')
add_bullets(doc, [
    'HangXe -> DongXe: một hãng có nhiều dòng xe',
    'DongXe -> PhienBanXe: một dòng xe có nhiều phiên bản',
    'TaiKhoan -> DonDatCoc: một tài khoản có nhiều đơn cọc',
    'PhienBanXe -> DonDatCoc: một phiên bản có thể nằm trong nhiều đơn cọc',
    'DonDatCoc -> HoaDonMuaXe: một đơn cọc có thể tạo một hóa đơn mua xe',
    'TaiKhoan -> GioHang: một tài khoản có nhiều dòng giỏ hàng'
])
add_paragraph(doc, '[CHÈN ẢNH: ERD hoặc ảnh sơ đồ quan hệ database đã vẽ bằng draw.io / dbdiagram]')
add_page_break(doc)

add_paragraph(doc, '7. Use case diagram và luồng nghiệp vụ', bold=True)
doc.add_heading('7. Use case diagram và luồng nghiệp vụ', level=1)
add_table(doc, ['Actor', 'Use case chính'], [
    ['Khách hàng', 'Xem xe, đặt lịch lái thử, đăng ký, đặt cọc, xem hồ sơ'],
    ['Admin', 'CRUD xe, CRUD người dùng, quản lý đơn cọc, nhận hóa đơn, theo dõi logs'],
    ['Quản lý showroom', 'Xem dashboard doanh thu, xem chi nhánh, quản lý lịch hẹn'],
    ['Hệ thống thanh toán', 'Tạo QR, xác minh webhook, cập nhật trạng thái đơn']
])
add_code_block(doc, [
    'graph TD',
    'A[Khách hàng] --> B[Xem xe]',
    'A --> C[Đặt cọc]',
    'A --> D[Đặt lịch lái thử]',
    'E[Admin] --> F[Quản lý xe]',
    'E --> G[Quản lý đơn cọc]',
    'H[Quản lý showroom] --> I[Dashboard doanh thu]'
])
add_page_break(doc)

add_paragraph(doc, '8. Class diagram gợi ý cho hệ thống', bold=True)
doc.add_heading('8. Class diagram gợi ý cho hệ thống', level=1)
add_code_block(doc, [
    'classDiagram',
    'class TaiKhoan {',
    '  +string TenDangNhap',
    '  +string VaiTro',
    '  +string Email',
    '}',
    'class HangXe { +int MaHang +string TenHang }',
    'class DongXe { +int MaDong +int MaHang +string TenDong }',
    'class PhienBanXe { +int MaPhienBan +int MaDong +long GiaNiemYet +int SoLuongTrongKho }',
    'class DonDatCoc { +int MaDonCoc +string MaKhachHang +int MaPhienBan +decimal SoTienCoc }',
    'class HoaDonMuaXe { +string MaHoaDon +int MaDonCoc +long TongTienPhaiTra }',
    'class GioHang { +int MaGioHang +string MaTaiKhoan +int MaPhienBan +int SoLuong }',
    'TaiKhoan --> DonDatCoc',
    'PhienBanXe --> DonDatCoc',
    'DonDatCoc --> HoaDonMuaXe',
    'TaiKhoan --> GioHang',
    'PhienBanXe --> GioHang'
])
add_page_break(doc)

add_paragraph(doc, '9. Triển khai Docker và SQL Server', bold=True)
doc.add_heading('9. Triển khai Docker và SQL Server', level=1)
add_paragraph(doc, 'Docker được dùng để containerize SQL Server. File docker-compose.yml định nghĩa service sqlserver với image mcr.microsoft.com/mssql/server:2022-latest, container_name sqlserver-carproject, port 1433 và volume sqlserver_data. Điều này giúp nhóm có thể khởi động database gần như tự động trên máy khác mà không cần cài SQL Server trực tiếp.')
add_bullets(doc, [
    'Khởi động: docker compose up -d',
    'Xem container: docker ps',
    'Xem log: docker logs sqlserver-carproject',
    'Dừng: docker compose down',
    'Migrate database: dotnet ef database update hoặc tự động qua startup'
])
add_paragraph(doc, '[CHÈN ẢNH: Docker Desktop đang chạy SQL Server và port mapping 1433:1433]')
add_page_break(doc)

add_paragraph(doc, '10. Kết nối và truy cập database cho nhóm qua Tailscale', bold=True)
doc.add_heading('10. Kết nối và truy cập database cho nhóm qua Tailscale', level=1)
add_paragraph(doc, 'Trong file cấu hình, hệ thống hiện dùng IP 100.108.48.1 để kết nối tới SQL Server. Nếu nhóm dùng Tailscale, người tham gia có thể kết nối tới cùng một tailnet, sau đó mở quyền truy cập SQL Server từ các máy khác.')
add_bullets(doc, [
    'Bước 1: cài Tailscale trên máy host và thành viên nhóm',
    'Bước 2: đăng nhập cùng account và join vào cùng tailnet',
    'Bước 3: mở port 1433 của SQL Server cho các IP ở trong tailnet',
    'Bước 4: dùng SQL Server Management Studio hay Azure Data Studio kết nối tới IP Tailscale hoặc IP nội bộ host'
])
add_paragraph(doc, '[CHÈN ẢNH: Màn hình Tailscale hoặc cấu hình mạng nội bộ]')
add_page_break(doc)

add_paragraph(doc, '11. Cloudflare và cấu hình domain mylxcar.online', bold=True)
doc.add_heading('11. Cloudflare và cấu hình domain mylxcar.online', level=1)
add_paragraph(doc, 'Nếu hệ thống triển khai lên internet, Cloudflare có thể dùng để làm reverse proxy, DNS và bảo mật cho domain mylxcar.online. Các bước cơ bản là tạo DNS A/AAAA record, bật proxy, cấu hình SSL/TLS Full hoặc Full Strict, và xác nhận reverse proxy đang chuyển tiếp tới ASP.NET Core.')
add_bullets(doc, [
    'DNS A record: mylxcar.online -> IP máy chủ',
    'DNS CNAME: www -> mylxcar.online',
    'SSL/TLS: bật Full hoặc Full Strict',
    'Cloudflare WAF/Rate limiting: khuyến nghị bật cho production',
    'Nginx/Traefik/Apache: chuyển tiếp traffic tới cổng 5001 hoặc 80 của ứng dụng'
])
add_paragraph(doc, '[CHÈN ẢNH: Cloudflare DNS / SSL / proxy configuration]')
add_page_break(doc)

add_paragraph(doc, '12. Bảo mật, xác thực và phân quyền', bold=True)
doc.add_heading('12. Bảo mật, xác thực và phân quyền', level=1)
add_paragraph(doc, 'Hệ thống dùng cookie session, JWT bearer và Google OAuth. Mỗi role có quyền khác nhau: Admin có toàn quyền CRUD; Quản lý được phép xem dashboard cho chi nhánh của mình; User chỉ có thể xem và đặt cọc. Điều này được xử lý trong middleware liên quan tới request path.')
add_bullets(doc, [
    'Cookie HttpOnly để giảm nguy cơ XSS',
    'JWT HMAC-SHA256 với secret key',
    'Authorization middleware kiểm soát /Admin/*, /Orders/*, /QuanLy/*',
    'Ghi log IP, browser, route, hành động',
    'Phân quyền theo role để tránh rò dữ liệu'
])
add_page_break(doc)

add_paragraph(doc, '13. Hướng dẫn chèn ảnh và biểu đồ vào báo cáo', bold=True)
doc.add_heading('13. Hướng dẫn chèn ảnh và biểu đồ vào báo cáo', level=1)
add_bullets(doc, [
    'Ảnh 1: giao diện landing page của hệ thống (đặt sau mục 2)',
    'Ảnh 2: ảnh database và kết nối SQL Server Management Studio (đặt sau mục 4)',
    'Ảnh 3: ERD / dbdiagram / draw.io (đặt sau mục 6)',
    'Ảnh 4: use case diagram (đặt sau mục 7)',
    'Ảnh 5: class diagram (đặt sau mục 8)',
    'Ảnh 6: Docker Desktop đang chạy SQL Server (đặt sau mục 9)',
    'Ảnh 7: màn hình Tailscale hoặc cấu hình mạng nội bộ (đặt sau mục 10)',
    'Ảnh 8: Cloudflare DNS / SSL / proxy configuration (đặt sau mục 11)',
    'Ảnh 9: dashboard Admin hoặc Quản lý showroom (đặt sau mục 12)'
])
add_page_break(doc)

add_paragraph(doc, '14. Mã Mermaid để dựng biểu đồ', bold=True)
doc.add_heading('14. Mã Mermaid để dựng biểu đồ', level=1)
add_paragraph(doc, 'Dưới đây là các đoạn mã để bạn copy sang Mermaid Live Editor hoặc PlantUML để tạo ảnh cho báo cáo.')
add_code_block(doc, [
    'flowchart TD',
    'A[Khách hàng] --> B[Xem xe]',
    'A --> C[Đặt cọc]',
    'A --> D[Đặt lịch lái thử]',
    'E[Admin] --> F[Quản lý xe]',
    'E --> G[Quản lý đơn cọc]',
    'H[Quản lý showroom] --> I[Dashboard doanh thu]'
])
add_paragraph(doc, '')
add_code_block(doc, [
    'classDiagram',
    'TaiKhoan --> GioHang',
    'DongXe --> PhienBanXe',
    'DonDatCoc --> HoaDonMuaXe'
])
add_paragraph(doc, '')
add_code_block(doc, [
    'graph LR',
    'Client[Browser] --> Cloudflare[Cloudflare DNS/SSL]',
    'Cloudflare --> App[ASP.NET Core App]',
    'App --> DB[SQL Server Docker]',
    'Team[Nhóm phát triển] --> Tailscale[Tailscale Network]',
    'Tailscale --> DB'
])
add_page_break(doc)

add_paragraph(doc, '15. Kết luận', bold=True)
doc.add_heading('15. Kết luận', level=1)
add_paragraph(doc, 'Hệ thống MyLxCar là một nền tảng web quản lý showroom xe hạng sang có đầy đủ chức năng bán hàng, quản trị, thanh toán và lưu trữ dữ liệu. Từ việc đọc toàn bộ code, có thể thấy hệ thống đang đi đúng hướng về kiến trúc, phân quyền, ORM, log và triển khai container. Điểm mạnh là có thể triển khai linh hoạt trên Docker, dùng Tailscale cho nhóm và Cloudflare cho domain. Điểm cần cải tiến tiếp là chuẩn hóa dữ liệu theo 3NF, tăng cường bảo mật, và tối ưu cấu hình production cho domain mylxcar.online.')
add_paragraph(doc, 'Đây là một báo cáo phù hợp để dùng cho đồ án hoặc buổi bảo vệ, vì nó vừa có tư duy hệ thống, vừa có sơ đồ, code mẫu, lưu ý triển khai thật và các vị trí cần chèn hình ảnh để tăng tính thuyết phục.')

doc.save(output_path)
print('Created', output_path)
