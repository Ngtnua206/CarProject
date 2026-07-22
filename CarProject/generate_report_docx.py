from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

output_path = r"D:\Code\Code\WebMVC\CarProject\BaoCao_HeThong_MyLxCar.docx"


def set_style(doc):
    styles = doc.styles
    style = styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    heading1 = styles['Heading 1']
    heading1.font.name = 'Arial'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = styles['Heading 2']
    heading2.font.name = 'Arial'
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = styles['Heading 3']
    heading3.font.name = 'Arial'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(3)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False, align=None, font_size=None, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
    for row_values in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_values):
            row_cells[i].text = value
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    return table


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = None


def add_image_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"[CHÈN ẢNH: {label}]")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
set_style(doc)

# Cover
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.add_run('BÁO CÁO HỆ THỐNG MYLXCAR')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Arial'
run.font.color.rgb = None

p2 = doc.add_paragraph()
p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run2 = p2.add_run('Website bán xe hạng sang, quản lý showroom, đơn cọc, lịch hẹn, thống kê doanh thu và tích hợp thanh toán')
run2.font.size = Pt(12)
run2.italic = True

doc.add_paragraph('Sinh viên: Nhóm phát triển MyLxCar')
doc.add_paragraph('Môn học: Hệ thống thông tin quản lý')
doc.add_paragraph('Ngày tạo: 22/07/2026')
doc.add_paragraph('Nền tảng: ASP.NET Core Razor Pages + SQL Server + Docker + Cloudflare + Tailscale')
add_page_break(doc)

# Section 1
add_heading(doc, '1. Tổng quan hệ thống', 1)
add_paragraph(doc, 'Hệ thống MyLxCar là nền tảng web bán hàng và quản trị showroom xe hạng sang. Trong code hiện tại, hệ thống gồm các thành phần chính: giao diện Razor Pages, service layer, Entity Framework Core, SQL Server chạy trong Docker, và middleware xác thực/ phân quyền.')
add_paragraph(doc, 'Từ việc đọc toàn bộ code, hệ thống hỗ trợ người dùng xem danh sách xe, xem chi tiết, so sánh xe, đặt lịch lái thử, đặt cọc, theo dõi giỏ hàng, đăng ký tài khoản, xác nhận email, thanh toán qua Sepay, và quản trị showroom cho Admin/Quản lý.')
add_paragraph(doc, 'Đây là hệ thống có thể triển khai cho cả môi trường nội bộ và môi trường công khai bằng cách đặt reverse proxy ở Cloudflare hoặc Nginx, đồng thời cho các thành viên trong nhóm truy cập database qua Tailscale.')
add_table(doc, ['Thành phần', 'Vai trò trong hệ thống'], [
    ['Frontend', 'Razor Pages + Bootstrap + Chart.js'],
    ['Backend', 'ASP.NET Core 10, middleware, service classes'],
    ['Database', 'SQL Server 2022 trên Docker'],
    ['Bảo mật', 'Session, JWT, Google OAuth, role-based authorization'],
    ['Thanh toán', 'Sepay QR + webhook'],
    ['Giám sát', 'Serilog, activity log, audit logs']
])
add_page_break(doc)

# Section 2
add_heading(doc, '2. Kiến trúc tổng thể và luồng xử lý', 1)
add_paragraph(doc, 'Ứng dụng được xây dựng theo kiến trúc 3 lớp: Presentation, Application/Service, Data Access. Presentation là Razor Pages và Controllers. Service layer chứa logic về giỏ hàng, email, QR thanh toán, JWT và log. Data layer là AppDbContext và các entity models. Mỗi request đi qua pipeline middleware rồi tới page model hoặc endpoint.')
add_paragraph(doc, 'Trong Program.cs, hệ thống đăng ký các dịch vụ: AddRazorPages, AddControllersWithViews, AddSession, AddAuthentication, AddAuthorization. Hệ thống cũng cấu hình Kestrel, giới hạn upload, Serilog và tự động migrate database khi chạy.')
add_bullets(doc, [
    'Luồng đăng nhập: request -> middleware -> authentication -> page handler -> service -> database',
    'Luồng đặt cọc: giỏ hàng -> tạo đơn cọc -> gọi Sepay -> lưu giao dịch -> cập nhật trạng thái',
    'Luồng quản trị: Admin/Quản lý -> dashboard -> CRUD -> ghi log hệ thống',
    'Luồng lịch hẹn lái thử: user -> form -> lưu vào LichHenLaiThu -> admin xem và xử lý'
])
add_image_placeholder(doc, 'Sơ đồ kiến trúc tổng thể (frontend, backend, database, Docker, Cloudflare, Tailscale)')
add_page_break(doc)

# Section 3
add_heading(doc, '3. Cấu trúc code và logic nghiệp vụ của hệ thống', 1)
add_paragraph(doc, 'Program.cs là điểm khởi đầu của ứng dụng. Ở đây, hệ thống đăng ký authentication, session, JWT, Google OAuth, DbContext, service scoped và middleware phân quyền theo đường dẫn. Đây là nơi quy định mọi route /Admin/* chỉ dành cho Admin, còn /Orders/* và /QuanLy/* yêu cầu đăng nhập.')
add_paragraph(doc, 'AppDbContext định nghĩa các DbSet và các mối quan hệ giữa bảng. Trong code, các bảng như HangXe, DongXe, PhienBanXe, TaiKhoan, DonDatCoc, HoaDonMuaXe, LichHenLaiThu, ChiNhanhShowroom, ChuongTrinhKhuyenMai, QuangCaoBanner, KenhTuVan, NhatKyHeThong, ThongKeTongHop_Boss và GioHang đều được ánh xạ với các khóa ngoại.')
add_table(doc, ['Module', 'Logic chính'], [
    ['CartService', 'Quản lý giỏ hàng, số lượng, tổng tiền cọc, kiểm tra tối thiểu 3 xe'],
    ['EmailService', 'Gửi email xác nhận đăng ký sử dụng SMTP'],
    ['SepayService', 'Tạo QR và verify webhook thanh toán'],
    ['DbInitializer', 'Seed dữ liệu ban đầu và tạo user AppReader'],
    ['ActivityLogService', 'Ghi log hành động người dùng, trình duyệt, IP, route']
])
add_paragraph(doc, 'Đặc biệt, CartService dùng session UserName để xác định người dùng hiện tại; mỗi giỏ hàng được liên kết với một tài khoản và một phiên bản xe. Đây là logic quan trọng để hỗ trợ đặt cọc và xác thực số lượng theo quy tắc tối thiểu 3 sản phẩm.')
add_page_break(doc)

# Section 4
add_heading(doc, '4. Cơ sở dữ liệu và mô hình dữ liệu', 1)
add_paragraph(doc, 'Cơ sở dữ liệu được thiết kế trên SQL Server 2022 chạy trong container Docker. Connection string hiện tại trong file appsettings.Development.json dùng Server=100.108.48.1,1433;Database=CarShopDb;User Id=sa;Password=Iumaioanhh@2024;TrustServerCertificate=True;Encrypt=False;.')
add_paragraph(doc, 'Mạng và kết nối giữa các máy trong nhóm có thể dùng Tailscale để tránh dùng public IP. Nếu một thành viên đã join vào tailnet, họ chỉ cần kết nối tới địa chỉ Tailscale của server hoặc IP nội bộ của máy hosting, sau đó mở cổng 1433 cho SQL Server.')
add_bullets(doc, [
    'Database Name: CarShopDb',
    'Container SQL: sqlserver-carproject',
    'Port nội bộ: 1433',
    'Phương thức kết nối: SQL Server Authentication',
    'Mục tiêu: nhóm có thể truy cập DB từ xa mà không cần public IP phơi lộ'
])
add_image_placeholder(doc, 'Màn hình kết nối SQL Server Management Studio / Azure Data Studio vào database')
add_page_break(doc)

# Section 5
add_heading(doc, '5. Phân tích chuẩn hóa dữ liệu 3NF', 1)
add_paragraph(doc, 'Mô hình dữ liệu hiện tại đã có hướng đi đúng đắn về phân tách thực thể và khóa ngoại, nhưng để đạt chuẩn 3NF một cách thỏa đáng hơn, cần lưu ý một số điểm.')
add_paragraph(doc, 'Đối với 3NF, mỗi bảng nên chứa dữ liệu chỉ về một thực thể và mọi thuộc tính phụ thuộc hoàn toàn vào khóa chính. Trong hệ thống hiện tại, bảng HangXe, DongXe, PhienBanXe, TaiKhoan, ChiNhanhShowroom, ChuongTrinhKhuyenMai là ở dạng khá tốt. Tuy nhiên, một số trường như HoTen, SoDienThoai, DiaChi, GhiChu trong DonDatCoc nên được kiểm tra lại vì chúng là thuộc tính mô tả khách hàng và có thể chuyển sang thực thể khách hàng hoặc bảng chi tiết khách hàng để tránh lặp dữ liệu.')
add_table(doc, ['Bảng', 'Đánh giá 3NF', 'Gợi ý cải tiến'], [
    ['HangXe', 'Đạt', 'Giữ nguyên'],
    ['DongXe', 'Đạt', 'Giữ nguyên và nối với HangXe'],
    ['PhienBanXe', 'Đạt', 'Giữ nguyên, phân loại theo dong xe'],
    ['TaiKhoan', 'Đạt', 'Giữ thông tin đăng nhập và vai trò'],
    ['DonDatCoc', 'Gần 3NF', 'Tách thông tin khách hàng sang bảng chi tiết khách hàng hoặc bảng KhachHang'],
    ['HoaDonMuaXe', 'Đạt', 'Không nên lưu trùng các trường khác ngoài thông tin hóa đơn'],
    ['LichHenLaiThu', 'Đạt', 'Có thể tách trạng thái sang bảng tham chiếu trạng thái'],
    ['GioHang', 'Đạt', 'Phụ thuộc vào tài khoản và phiên bản xe']
])
add_paragraph(doc, 'Kết luận: mô hình hiện tại có thể được xem là gần 3NF, phù hợp cho việc học và triển khai ban đầu; nếu muốn chuẩn hóa hoàn chỉnh, nên tách các thuộc tính khách hàng khỏi đơn cọc và dùng khóa ngoại đến bảng khách hàng/chitiet khách hàng.')
add_page_break(doc)

# Section 6
add_heading(doc, '6. Mối quan hệ giữa các bảng và ERD', 1)
add_paragraph(doc, 'Hệ thống dùng nhiều bảng liên kết với nhau bằng khóa ngoại. Quan hệ chính có thể được mô tả như sau: HangXe 1-n DongXe; DongXe 1-n PhienBanXe; TaiKhoan 1-n DonDatCoc; PhienBanXe 1-n DonDatCoc; DonDatCoc 1-1 HoaDonMuaXe; TaiKhoan 1-n LichHenLaiThu; DongXe 1-n LichHenLaiThu; ChiNhanhShowroom 1-n LichHenLaiThu.')
add_paragraph(doc, 'Ngoài ra, GioHang là bảng trung gian giữa TaiKhoan và PhienBanXe, cho phép người dùng lưu nhiều xe trong giỏ hàng trước khi đặt cọc. Bảng NhatKyHeThong ghi các hành động của user và admin, và được nối tới TaiKhoan.')
add_bullets(doc, [
    'HangXe -> DongXe: một hãng có nhiều dòng xe',
    'DongXe -> PhienBanXe: một dòng xe có nhiều phiên bản',
    'TaiKhoan -> DonDatCoc: một tài khoản có nhiều đơn cọc',
    'PhienBanXe -> DonDatCoc: một phiên bản có thể nằm trong nhiều đơn cọc',
    'DonDatCoc -> HoaDonMuaXe: một đơn cọc có thể tạo một hóa đơn mua xe',
    'TaiKhoan -> GioHang: một tài khoản có nhiều dòng giỏ hàng'
])
add_image_placeholder(doc, 'ERD hoặc ảnh sơ đồ quan hệ database đã vẽ bằng draw.io / dbdiagram')
add_page_break(doc)

# Section 7
add_heading(doc, '7. Use case diagram và luồng nghiệp vụ', 1)
add_paragraph(doc, 'Các actor chính gồm: Khách hàng, Admin, Quản lý showroom, hệ thống thanh toán, hệ thống email.')
add_table(doc, ['Actor', 'Use case chính'], [
    ['Khách hàng', 'Xem xe, đặt lịch lái thử, đăng ký, đặt cọc, xem hồ sơ'],
    ['Admin', 'CRUD xe, CRUD người dùng, quản lý đơn cọc, nhận hóa đơn, theo dõi logs'],
    ['Quản lý showroom', 'Xem dashboard doanh thu, xem chi nhánh, quản lý lịch hẹn'],
    ['Hệ thống thanh toán', 'Tạo QR, xác minh webhook, cập nhật trạng thái đơn']
])
add_paragraph(doc, 'Đây là một hệ thống có hai nhóm người dùng chính: người dùng công khai và quản trị. Mỗi use case đều liên kết tới một module code riêng trong Razor Pages hoặc service layer.')
add_code_block(doc, [
    'graph TD',
    'A[Khách hàng] --> B[Xem xe]',
    'A --> C[Đặt lịch lái thử]',
    'A --> D[Đặt cọc]',
    'A --> E[Quản lý hồ sơ]',
    'F[Admin] --> G[CRUD xe]',
    'F --> H[Quản lý đơn cọc]',
    'I[Quản lý showroom] --> J[Dashboard doanh thu]'
])
add_page_break(doc)

# Section 8
add_heading(doc, '8. Class diagram gợi ý cho hệ thống', 1)
add_paragraph(doc, 'Nếu cần vẽ biểu đồ lớp, có thể nhóm theo domain model và service. Nền tảng hiện tại có các entity class như HangXe, DongXe, PhienBanXe, TaiKhoan, DonDatCoc, HoaDonMuaXe, GioHang, LichHenLaiThu, ChiNhanhShowroom, QuangCaoBanner và service như CartService, SepayService, EmailService.')
add_code_block(doc, [
    'classDiagram',
    'class TaiKhoan {',
    '  +string TenDangNhap',
    '  +string VaiTro',
    '  +string Email',
    '}',
    'class HangXe { +int MaHang +string TenHang }',
    'class DongXe { +int MaDong +int MaHang +string TenDong }',
    'class PhienBanXe { +int MaPhienBan +int MaDong +long GiaNiemYet +int SoLuongTrongKho }',
    'class DonDatCoc { +int MaDonCoc +string MaKhachHang +int MaPhienBan +decimal SoTienCoc }',
    'class HoaDonMuaXe { +string MaHoaDon +int MaDonCoc +long TongTienPhaiTra }',
    'class GioHang { +int MaGioHang +string MaTaiKhoan +int MaPhienBan +int SoLuong }',
    'TaiKhoan --> DonDatCoc',
    'PhienBanXe --> DonDatCoc',
    'DonDatCoc --> HoaDonMuaXe',
    'TaiKhoan --> GioHang',
    'PhienBanXe --> GioHang'
])
add_page_break(doc)

# Section 9
add_heading(doc, '9. Triển khai Docker và SQL Server', 1)
add_paragraph(doc, 'Docker được dùng để containerize SQL Server. File docker-compose.yml định nghĩa service sqlserver với image mcr.microsoft.com/mssql/server:2022-latest, container_name sqlserver-carproject, port 1433 và volume sqlserver_data. Điều này giúp nhóm có thể khởi động database gần như tự động trên máy khác mà không cần cài SQL Server trực tiếp.')
add_paragraph(doc, 'Khi chạy ứng dụng, AppDbContext dùng EF Core migrate để tạo bảng và seed dữ liệu ban đầu. DbInitializer còn tạo một login AppReader cho mục đích đọc dữ liệu và nhắm vào quyền hạn hợp lý.')
add_bullets(doc, [
    'Khởi động: docker compose up -d',
    'Xem container: docker ps',
    'Xem log: docker logs sqlserver-carproject',
    'Dừng: docker compose down',
    'Migrate database: dotnet ef database update hoặc tự động qua startup'
])
add_image_placeholder(doc, 'Ảnh Docker Desktop đang chạy container SQL Server và port mapping 1433:1433')
add_page_break(doc)

# Section 10
add_heading(doc, '10. Kết nối và truy cập database cho nhóm qua Tailscale', 1)
add_paragraph(doc, 'Trong file cấu hình, hệ thống hiện dùng IP 100.108.48.1 để kết nối tới SQL Server. Đây có thể là địa chỉ nội bộ hoặc Tailscale IP của máy hosting. Nếu nhóm dùng Tailscale, người tham gia có thể kết nối tới cùng một tailnet, sau đó mở quyền truy cập SQL Server từ các máy khác.')
add_paragraph(doc, 'Để làm việc thuận tiện, nên chuẩn hóa một trong các phương án sau: (1) dùng Tailscale IP của server, (2) dùng hostname nội bộ nếu DNS nội bộ sẵn có, hoặc (3) dùng public IP + Cloudflare Tunnel nếu cần public access.')
add_bullets(doc, [
    'Bước 1: cài Tailscale trên máy host và thành viên nhóm',
    'Bước 2: đăng nhập cùng account và join vào cùng tailnet',
    'Bước 3: mở port 1433 của SQL Server cho các IP ở trong tailnet',
    'Bước 4: dùng SQL Server Management Studio hay Azure Data Studio kết nối tới IP Tailscale hoặc IP nội bộ host'
])
add_paragraph(doc, 'Lưu ý bảo mật: không nên để 1433 mở công khai trên Internet; chỉ mở cho các máy có trong tailnet hoặc cho các IP xác định. Nếu cần truy cập từ ngoài, nên dùng VPN hoặc Cloudflare Tunnel thay vì mở cổng trực tiếp.')
add_page_break(doc)

# Section 11
add_heading(doc, '11. Cloudflare và cấu hình domain mylxcar.online', 1)
add_paragraph(doc, 'Nếu hệ thống triển khai lên internet, Cloudflare có thể dùng để làm reverse proxy, DNS và bảo mật cho domain mylxcar.online. Nhiệm vụ của Cloudflare là đưa lưu lượng từ domain vào máy chủ hosting, đồng thời cung cấp TLS/SSL và bảo vệ khỏi các cuộc tấn công phổ biến.')
add_paragraph(doc, 'Các bước cơ bản là: tạo DNS A/AAAA record cho mylxcar.online hoặc www để trỏ tới IP máy chủ, bật proxy (orange cloud), cấu hình SSL/TLS Full hoặc Full Strict, và xác nhận rằng reverse proxy hoặc Nginx ở máy chủ đang lắng nghe cổng 80/443 và chuyển tiếp tới ASP.NET Core.')
add_bullets(doc, [
    'DNS A record: mylxcar.online -> IP máy chủ',
    'DNS CNAME: www -> mylxcar.online',
    'SSL/TLS: bật Full hoặc Full Strict',
    'Cloudflare WAF/Rate limiting: khuyến nghị bật cho môi trường production',
    'Nginx/Traefik/Apache: chuyển tiếp traffic tới cổng 5001 hoặc 80 của ứng dụng']
])
add_paragraph(doc, 'Với hệ thống hiện tại, vì ứng dụng chạy trên Kestrel và có thể mở cổng 5001, bạn có thể đặt reverse proxy phía trước để Cloudflare chỉ công khai 443 và không trực tiếp mở ứng dụng ra internet.')
add_page_break(doc)

# Section 12
add_heading(doc, '12. Bảo mật, xác thực và phân quyền', 1)
add_paragraph(doc, 'Hệ thống dùng cookie session, JWT bearer và Google OAuth. Mỗi role có quyền khác nhau: Admin có toàn quyền CRUD; Quản lý được phép xem dashboard cho chi nhánh của mình; User chỉ có thể xem và đặt cọc. Điều này được xử lý trong middleware liên quan tới request path.')
add_paragraph(doc, 'Ngoài ra, sự kiện đăng nhập, đổ lỗi, upload avatar, thanh toán và thao tác admin đều được ghi log bằng Serilog và trong bảng NhatKyHeThong. Đây là một lợi thế quan trọng cho việc kiểm tra truy cập, phát hiện lỗi và trao đổi với nhóm.')
add_bullets(doc, [
    'Cookie HttpOnly để giảm nguy cơ XSS',
    'JWT HMAC-SHA256 với secret key',
    'Authorization middleware kiểm soát /Admin/*, /Orders/*, /QuanLy/*',
    'Ghi log IP, browser, route, hành động',
    'Phân quyền theo role để tránh rò dữ liệu']
])
add_page_break(doc)

# Section 13
add_heading(doc, '13. Các chức năng chính và nhận diện module', 1)
add_table(doc, ['Chức năng', 'Module liên quan'], [
    ['Trang chủ', 'Pages/Index'],
    ['Danh sách xe', 'Pages/Cars'],
    ['Chi tiết xe', 'Pages/Details'],
    ['So sánh xe', 'Pages/Compare'],
    ['Đặt lịch lái thử', 'Pages/TestDrive'],
    ['Đặt cọc và checkout', 'Pages/Orders/Cart/Checkout'],
    ['Quản lý hồ sơ', 'Pages/Profile'],
    ['Quản trị dashboard', 'Pages/Admin/Index / Pages/QuanLy/Dashboard']
])
add_paragraph(doc, 'Mỗi module có thể được mở rộng độc lập, vì cấu trúc Razor Pages và service layer giúp tách biệt UI và logic. Trong tương lai, nếu cần mở rộng thành microservice, việc tách PayService và EmailService là hợp lý.')
add_page_break(doc)

# Section 14
add_heading(doc, '14. Hướng dẫn chèn ảnh vào báo cáo', 1)
add_paragraph(doc, 'Để báo cáo đẹp và chuyên nghiệp hơn, nên chèn các ảnh sau ở các vị trí sau:')
add_bullets(doc, [
    'Ảnh 1: giao diện landing page của hệ thống (đặt sau mục 2. Kiến trúc tổng thể)',
    'Ảnh 2: ảnh database và kết nối SQL Server Management Studio (đặt sau mục 4. Cơ sở dữ liệu)',
    'Ảnh 3: ERD / dbdiagram / draw.io (đặt sau mục 6. Mối quan hệ giữa các bảng)',
    'Ảnh 4: use case diagram (đặt sau mục 7. Use case diagram)',
    'Ảnh 5: class diagram (đặt sau mục 8. Class diagram)',
    'Ảnh 6: Docker Desktop đang chạy SQL Server (đặt sau mục 9. Triển khai Docker)',
    'Ảnh 7: màn hình Tailscale hoặc cấu hình mạng nội bộ (đặt sau mục 10)',
    'Ảnh 8: Cloudflare DNS / SSL / proxy configuration (đặt sau mục 11)',
    'Ảnh 9: dashboard Admin hoặc Quản lý showroom (đặt sau mục 13)'
])
add_paragraph(doc, 'Nếu bạn cần, có thể chụp màn hình từ hệ thống đã chạy và dán vào đúng các vị trí trên. Với Word, bạn chỉ cần chọn Insert -> Pictures -> This Device ở đúng chỗ có placeholder.')
add_page_break(doc)

# Section 15
add_heading(doc, '15. Mã Mermaid và PlantUML để dựng biểu đồ', 1)
add_paragraph(doc, 'Dưới đây là các đoạn mã để bạn copy sang Mermaid Live Editor hoặc PlantUML để tạo ảnh cho báo cáo.')
add_heading(doc, '15.1. Mermaid – Use Case', 2)
add_code_block(doc, [
    'flowchart TD',
    'A[Khách hàng] --> B[Xem xe]',
    'A --> C[Đặt cọc]',
    'A --> D[Đặt lịch lái thử]',
    'E[Admin] --> F[Quản lý xe]',
    'E --> G[Quản lý đơn cọc]',
    'H[Quản lý showroom] --> I[Dashboard doanh thu]'
])
add_heading(doc, '15.2. Mermaid – Class Diagram', 2)
add_code_block(doc, [
    'classDiagram',
    'TaiKhoan <|-- Admin',
    'TaiKhoan <|-- User',
    'TaiKhoan --> GioHang',
    'DongXe --> PhienBanXe',
    'DonDatCoc --> HoaDonMuaXe'
])
add_heading(doc, '15.3. Mermaid – Deployment', 2)
add_code_block(doc, [
    'graph LR',
    'Client[Browser] --> Cloudflare[Cloudflare DNS/SSL]',
    'Cloudflare --> App[ASP.NET Core App]',
    'App --> DB[SQL Server Docker]',
    'Team[Nhóm phát triển] --> Tailscale[Tailscale Network]',
    'Tailscale --> DB'
])
add_page_break(doc)

# Section 16
add_heading(doc, '16. Kết luận', 1)
add_paragraph(doc, 'Hệ thống MyLxCar là một nền tảng web quản lý showroom xe hạng sang có đầy đủ chức năng bán hàng, quản trị, thanh toán và lưu trữ dữ liệu. Từ việc đọc toàn bộ code, có thể thấy hệ thống đang đi đúng hướng về kiến trúc, phân quyền, ORM, log và triển khai container. Điểm mạnh là có thể triển khai linh hoạt trên Docker, dùng Tailscale cho nhóm và Cloudflare cho domain. Điểm cần cải tiến tiếp là chuẩn hóa dữ liệu theo 3NF, tăng cường bảo mật, và tối ưu cấu hình production cho domain mylxcar.online.')
add_paragraph(doc, 'Đây là một báo cáo phù hợp để dùng cho đồ án hoặc buổi bảo vệ, vì nó vừa có tư duy hệ thống, vừa có sơ đồ, code mẫu, lưu ý triển khai thật và các vị trí cần chèn hình ảnh để tăng tính thuyết phục.')

# Save
try:
    doc.save(output_path)
    print('Created', output_path)
except Exception as e:
    print('Error', e)
